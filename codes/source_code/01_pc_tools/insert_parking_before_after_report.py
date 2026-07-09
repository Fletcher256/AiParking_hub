from pathlib import Path
from PIL import Image, ImageOps, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

SRC = Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版_补实物照片.docx")
if not SRC.exists():
    SRC = Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版.docx")
OUT = Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版_补实物照片_补泊车前后对比.docx")
BEFORE = Path(r"C:/Users/Cheng/Downloads/IMG_6405.JPG")
AFTER = Path(r"C:/Users/Cheng/Downloads/IMG_6406.JPG")
COMPARE = Path(r"D:/parking_board_agent/docs/parking_before_after_compare.jpg")

for p in [SRC, BEFORE, AFTER]:
    if not p.exists():
        raise FileNotFoundError(p)


def get_font(size=34, bold=False):
    candidates = [
        r"C:/Windows/Fonts/msyhbd.ttc" if bold else r"C:/Windows/Fonts/msyh.ttc",
        r"C:/Windows/Fonts/simhei.ttf",
        r"C:/Windows/Fonts/simsun.ttc",
        r"C:/Windows/Fonts/arial.ttf",
    ]
    for c in candidates:
        try:
            return ImageFont.truetype(c, size)
        except Exception:
            pass
    return ImageFont.load_default()

FONT_TITLE = get_font(48, True)
FONT_SUB = get_font(28, False)
FONT_LABEL = get_font(39, True)
FONT_SMALL = get_font(25, False)
FONT_ARROW = get_font(32, True)


def fit_contain(img, size):
    img = ImageOps.exif_transpose(img).convert("RGB")
    w, h = img.size
    tw, th = size
    scale = min(tw / w, th / h)
    nw, nh = int(w * scale + 0.5), int(h * scale + 0.5)
    resized = img.resize((nw, nh), Image.Resampling.LANCZOS)
    bg = Image.new("RGB", size, (250, 250, 250))
    bg.paste(resized, ((tw - nw) // 2, (th - nh) // 2))
    return bg


def draw_panel(canvas, x, y, w, h, img_path, label, note, accent):
    d = ImageDraw.Draw(canvas)
    d.rounded_rectangle((x, y, x+w, y+h), radius=22, fill=(255,255,255), outline=accent, width=5)
    strip_h = 86
    d.rounded_rectangle((x, y, x+w, y+strip_h), radius=22, fill=accent, outline=accent)
    # square off lower part of strip so radius only appears top
    d.rectangle((x, y+strip_h-22, x+w, y+strip_h), fill=accent)
    bbox = d.textbbox((0,0), label, font=FONT_LABEL)
    d.text((x + (w - (bbox[2]-bbox[0]))//2, y+18), label, fill=(255,255,255), font=FONT_LABEL)
    img_box_h = h - strip_h - 80
    im = fit_contain(Image.open(img_path), (w-42, img_box_h))
    canvas.paste(im, (x+21, y+strip_h+18))
    # Note
    nb = d.textbbox((0,0), note, font=FONT_SMALL)
    d.text((x + (w-(nb[2]-nb[0]))//2, y+h-50), note, fill=(57,70,86), font=FONT_SMALL)


def draw_arrow(canvas, x1, y, x2, color=(22, 118, 180)):
    d = ImageDraw.Draw(canvas)
    d.line((x1, y, x2, y), fill=color, width=10)
    head = [(x2, y), (x2-42, y-28), (x2-42, y+28)]
    d.polygon(head, fill=color)
    text = "闭环感知-决策-执行"
    bbox = d.textbbox((0,0), text, font=FONT_ARROW)
    d.rounded_rectangle((x1+18, y-70, x2-18, y-26), radius=16, fill=(235, 246, 255), outline=color, width=2)
    d.text((x1 + (x2-x1-(bbox[2]-bbox[0]))//2, y-70), text, fill=color, font=FONT_ARROW)


def make_compare():
    COMPARE.parent.mkdir(parents=True, exist_ok=True)
    W, H = 2400, 1540
    canvas = Image.new("RGB", (W,H), (255,255,255))
    d = ImageDraw.Draw(canvas)
    margin = 70
    d.text((margin, 32), "泊车前后对比图", fill=(22,44,72), font=FONT_TITLE)
    d.text((margin, 88), "同一车位区域下，展示车辆由车位外侧初始姿态进入目标车位后的终态效果", fill=(82,96,116), font=FONT_SUB)
    panel_y = 150
    panel_h = 1300
    panel_w = 1040
    left_x = margin
    right_x = W - margin - panel_w
    draw_panel(canvas, left_x, panel_y, panel_w, panel_h, BEFORE, "泊车前", "车辆位于车位外侧，存在初始横向/航向偏差", (196, 88, 77))
    draw_panel(canvas, right_x, panel_y, panel_w, panel_h, AFTER, "泊车后", "车辆进入黄色胶带标识的车位区域", (41, 137, 92))
    draw_arrow(canvas, left_x+panel_w+16, panel_y+panel_h//2, right_x-16)
    d.text((margin, H-54), "注：第一张为泊车前，第二张为泊车后；照片用于支撑实车闭环演示与终态效果说明。", fill=(82,96,116), font=FONT_SMALL)
    canvas.save(COMPARE, quality=92, optimize=True)


def set_run_font(run, name="微软雅黑", size=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_text(paragraph, text, size=None, bold=None, italic=None):
    for r in paragraph.runs:
        r.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    set_run_font(run, size=size, bold=bold, italic=italic)


def move_after(new_para, cursor_para):
    cursor_para._p.addnext(new_para._p)
    return new_para


def add_before_after_after(doc, after_para):
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(6)
    intro.paragraph_format.space_after = Pt(4)
    intro.paragraph_format.keep_with_next = True
    r = intro.add_run("图证补充：泊车前后照片展示了车辆从车位外侧初始姿态，经视觉闭环决策和底盘执行后进入目标车位，可作为实车闭环效果的直观证据。")
    set_run_font(r, size=10.5)
    cursor = move_after(intro, after_para)

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(3)
    pic_p.paragraph_format.space_after = Pt(2)
    pic_p.paragraph_format.keep_with_next = True
    run = pic_p.add_run()
    run.add_picture(str(COMPARE), width=Inches(6.25))
    for drawing in run._element.xpath(".//w:drawing"):
        doc_pr = drawing.xpath(".//wp:docPr")
        if doc_pr:
            doc_pr[0].set("name", "泊车前后对比图")
            doc_pr[0].set("descr", "左侧为泊车前车辆位于车位外侧，右侧为泊车后车辆进入黄色胶带标识的目标车位区域。")
    cursor = move_after(pic_p, cursor)

    cap = doc.add_paragraph()
    try:
        cap.style = doc.styles["Caption"]
    except Exception:
        cap.style = doc.styles["Normal"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(9)
    set_paragraph_text(cap, "图3 泊车前后对比：车辆从车位外侧初始姿态出发，经视觉闭环控制后进入黄色胶带标识的目标车位区域。", size=9, italic=True)
    move_after(cap, cursor)


make_compare()

doc = Document(str(SRC))

# Strengthen verification table evidence if present.
for table in doc.tables:
    for row in table.rows:
        joined = " | ".join(cell.text for cell in row.cells)
        if "实车闭环" in joined and len(row.cells) >= 4:
            p = row.cells[3].paragraphs[0]
            text = p.text.strip()
            if "泊车前后对比" not in text:
                text = text.rstrip("。") + "；泊车前后对比见图3。"
            set_paragraph_text(p, text, size=9)

insert_after = None
# Prefer inserting right after the verification-result table caption.
for p in doc.paragraphs:
    if "表9 验证结果汇总" in p.text:
        insert_after = p
        break
if insert_after is None:
    # Fallback after 3.3 heading body if caption not found.
    for p in doc.paragraphs:
        if p.text.strip() == "3.3 验证结果":
            insert_after = p
            break
if insert_after is None:
    raise RuntimeError("未找到照片插入位置：表9 验证结果汇总 / 3.3 验证结果")

add_before_after_after(doc, insert_after)

cp = doc.core_properties
cp.title = "基于海鸥派 SS928 与 STM32 底盘的视觉闭环自主泊车系统"
cp.subject = "竞赛作品报告"
cp.author = ""
cp.comments = ""
cp.keywords = ""

doc.save(str(OUT))
print(f"SRC={SRC}")
print(f"OUT={OUT}")
print(f"COMPARE={COMPARE}")
