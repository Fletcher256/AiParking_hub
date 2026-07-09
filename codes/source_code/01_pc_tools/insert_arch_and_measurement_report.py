from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版_补实物照片_补泊车前后对比_补YOLO识别效果.docx")
OUT = Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版_最终补图表版.docx")
ARCH_IMG = Path(r"D:/parking_board_agent/docs/project_design_flowchart_preview.png")

if not SRC.exists():
    raise FileNotFoundError(SRC)
if not ARCH_IMG.exists():
    raise FileNotFoundError(ARCH_IMG)


def set_run_font(run, name="微软雅黑", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_text(paragraph, text, size=None, bold=None, italic=None):
    for r in paragraph.runs:
        r.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    set_run_font(run, size=size, bold=bold, italic=italic)


def move_after_para(new_para, cursor_para):
    cursor_para._p.addnext(new_para._p)
    return new_para


def move_table_after(table, cursor_para):
    cursor_para._p.addnext(table._tbl)
    return table


def delete_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def format_table(table):
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if ri == 0:
                shade_cell(cell, 'D9EAF7')
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=8.5 if ri else 9, bold=(ri == 0))


def add_architecture_block(doc, after_para):
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(6)
    intro.paragraph_format.space_after = Pt(4)
    intro.paragraph_format.keep_with_next = True
    r = intro.add_run("图证补充：系统整体架构图展示了从摄像头输入、YOLO 感知、位姿估计、决策规划到 STM32 底盘执行和日志复盘的完整闭环链路。")
    set_run_font(r, size=10.5)
    cursor = move_after_para(intro, after_para)

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(3)
    pic_p.paragraph_format.space_after = Pt(2)
    pic_p.paragraph_format.keep_with_next = True
    run = pic_p.add_run()
    run.add_picture(str(ARCH_IMG), width=Inches(6.25))
    for drawing in run._element.xpath(".//w:drawing"):
        doc_pr = drawing.xpath(".//wp:docPr")
        if doc_pr:
            doc_pr[0].set("name", "系统整体架构图")
            doc_pr[0].set("descr", "视觉闭环自主泊车系统整体架构图，展示摄像头、YOLO、位姿估计、决策规划、STM32执行和日志复盘链路。")
    cursor = move_after_para(pic_p, cursor)

    cap = doc.add_paragraph()
    try:
        cap.style = doc.styles["Caption"]
    except Exception:
        cap.style = doc.styles["Normal"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(9)
    set_paragraph_text(cap, "图2 系统整体架构图：系统由感知、位姿估计、决策控制、底盘执行和复盘记录组成，形成板端闭环控制链路。", size=9, italic=True)
    move_after_para(cap, cursor)


def add_measurement_table(doc, after_para):
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(6)
    intro.paragraph_format.space_after = Pt(4)
    intro.paragraph_format.keep_with_next = True
    r = intro.add_run("为便于现场或赛前复核终态精度，预留外部尺量记录表；其中“人工尺量/现场记录”列留空，由后续实测后填写。")
    set_run_font(r, size=10.5)
    cursor = move_after_para(intro, after_para)

    table = doc.add_table(rows=1, cols=4)
    headers = ["记录项", "当前已知/日志填写", "人工尺量/现场记录", "说明"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    rows = [
        ["演示来源", "2026-07-04 line_follow 实车录制流程", "—", "与原始 H264 视频、JSONL 控制日志对应。"],
        ["场景状态", "黄色胶带标识车位；车辆由车位外侧进入车位区域", "", "可结合图4泊车前后对比复核。"],
        ["动作过程", "12 步闭环动作；含一次前进调整段（shuffle）", "—", "每步执行后停止并重新观测/规划。"],
        ["日志估计终态", "y≈2.2 cm；lat≈-1.0 cm；head≈6.2°", "—", "来自控制日志估计结果。"],
        ["终态横向偏差", "待外部尺量复核", "", "填写车体中心或约定参考点到目标中心线的横向误差。"],
        ["终态航向角误差", "待外部尺量复核", "", "填写车体纵轴与车位目标方向的夹角。"],
        ["是否压线/越界", "待现场确认", "", "填写“否/是”，如有越界说明位置。"],
        ["人工复核结论", "待现场确认", "", "填写通过/需调整及简要原因。"],
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
    format_table(table)
    move_table_after(table, cursor)

    cap = doc.add_paragraph()
    try:
        cap.style = doc.styles["Caption"]
    except Exception:
        cap.style = doc.styles["Normal"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(9)
    set_paragraph_text(cap, "表11 实车终态外部尺量记录表（人工尺量项待填写）", size=9, italic=True)
    # Put caption after the moved table XML.
    table._tbl.addnext(cap._p)


doc = Document(str(SRC))

# 1) Remove appendix entries from the manual TOC/front matter because user does not want appendices.
for p in list(doc.paragraphs):
    t = p.text.strip()
    if t in {"附录A 竞赛展示与支撑材料清单", "附录B 现场风险与应对预案"}:
        delete_paragraph(p)

# 2) Insert architecture figure after the 2.1 explanatory paragraph.
insert_after = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "2.1 整体介绍":
        for q in doc.paragraphs[i+1:]:
            if q.text.strip() and not q.text.strip().startswith("表"):
                insert_after = q
                break
        break
if insert_after is None:
    raise RuntimeError("未找到 2.1 整体介绍后的插入位置")
add_architecture_block(doc, insert_after)

# 3) Renumber existing figure captions after inserting new Figure 2.
caption_replacements = {
    "图2 整车多角度实物照片": "图3 整车多角度实物照片",
    "图3 泊车前后对比": "图4 泊车前后对比",
    "图4 YOLO 车位识别效果": "图5 YOLO 车位识别效果",
}
for p in doc.paragraphs:
    txt = p.text
    for old, new in caption_replacements.items():
        if txt.startswith(old):
            set_paragraph_text(p, txt.replace(old, new, 1), size=9, italic=True)

# Update in-text/table references affected by figure renumbering.
for p in doc.paragraphs:
    txt = p.text
    new = txt.replace("泊车前后对比见图3", "泊车前后对比见图4").replace("识别效果见图4", "识别效果见图5")
    if new != txt:
        set_paragraph_text(p, new)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                txt = p.text
                new = txt.replace("泊车前后对比见图3", "泊车前后对比见图4").replace("识别效果见图4", "识别效果见图5")
                if new != txt:
                    # Keep table font small.
                    set_paragraph_text(p, new, size=9)

# 4) Insert external measurement table after Table 10 caption.
measure_after = None
for p in doc.paragraphs:
    if p.text.strip().startswith("表10 实车演示与量化验证方案"):
        measure_after = p
        break
if measure_after is None:
    raise RuntimeError("未找到 表10 插入位置")
add_measurement_table(doc, measure_after)

# 5) Neutral metadata; do not add reference [5].
cp = doc.core_properties
cp.title = "基于海鸥派 SS928 与 STM32 底盘的视觉闭环自主泊车系统"
cp.subject = "竞赛作品报告"
cp.author = ""
cp.comments = ""
cp.keywords = ""

doc.save(str(OUT))
print(f"OUT={OUT}")
