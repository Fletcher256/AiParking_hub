from pathlib import Path
from PIL import Image, ImageOps, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

SRC = Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版.docx")
OUT = Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版_补实物照片.docx")
COLLAGE = Path(r"D:/parking_board_agent/docs/vehicle_multiview_photos.jpg")
PHOTOS = [
    (Path(r"C:/Users/Cheng/Downloads/IMG_6399.JPG"), "a 正面：电控层与线束"),
    (Path(r"C:/Users/Cheng/Downloads/IMG_6401.JPG"), "b 左前斜上：SS928/STM32 分层安装"),
    (Path(r"C:/Users/Cheng/Downloads/IMG_6398.JPG"), "c 俯视：主控、接口与车体布局"),
    (Path(r"C:/Users/Cheng/Downloads/IMG_6403.JPG"), "d 侧视：底盘、电池与支撑结构"),
    (Path(r"C:/Users/Cheng/Downloads/IMG_6402.JPG"), "e 后部：摄像头与执行机构布置"),
]

for p, _ in PHOTOS:
    if not p.exists():
        raise FileNotFoundError(p)
if not SRC.exists():
    raise FileNotFoundError(SRC)


def get_font(size=34, bold=False):
    candidates = [
        r"C:/Windows/Fonts/msyhbd.ttc" if bold else r"C:/Windows/Fonts/msyh.ttc",
        r"C:/Windows/Fonts/simhei.ttf",
        r"C:/Windows/Fonts/simsun.ttc",
        r"C:/Windows/Fonts/arial.ttf",
    ]
    for c in candidates:
        try:
            return ImageFont.truetype(c, size)
        except Exception:
            pass
    return ImageFont.load_default()

FONT_LABEL = get_font(34, bold=True)
FONT_TITLE = get_font(46, bold=True)
FONT_NOTE = get_font(26, bold=False)


def cover_fit(img, size):
    # Crop-to-fill, centered, after EXIF rotation normalization.
    img = ImageOps.exif_transpose(img).convert("RGB")
    w, h = img.size
    tw, th = size
    scale = max(tw / w, th / h)
    nw, nh = int(w * scale + 0.5), int(h * scale + 0.5)
    img = img.resize((nw, nh), Image.Resampling.LANCZOS)
    left = max(0, (nw - tw) // 2)
    top = max(0, (nh - th) // 2)
    return img.crop((left, top, left + tw, top + th))


def draw_cell(canvas, xy, size, photo_path, label):
    x, y = xy
    w, h = size
    label_h = 54
    im = cover_fit(Image.open(photo_path), (w, h - label_h))
    canvas.paste(im, (x, y))
    d = ImageDraw.Draw(canvas)
    d.rectangle((x, y + h - label_h, x + w, y + h), fill=(245, 248, 252))
    d.rectangle((x, y, x + w, y + h), outline=(185, 195, 210), width=3)
    # Center label, fallback to left if too long.
    bbox = d.textbbox((0, 0), label, font=FONT_LABEL)
    text_w = bbox[2] - bbox[0]
    tx = x + max(18, (w - text_w) // 2)
    d.text((tx, y + h - label_h + 8), label, fill=(22, 44, 72), font=FONT_LABEL)


def make_collage():
    COLLAGE.parent.mkdir(parents=True, exist_ok=True)
    W, H = 2400, 1550
    margin = 70
    gap = 34
    title_h = 105
    foot_h = 70
    top_h = 610
    bot_h = 675
    top_w = (W - 2 * margin - 2 * gap) // 3
    bot_w = (W - 2 * margin - gap) // 2
    canvas = Image.new("RGB", (W, H), (255, 255, 255))
    d = ImageDraw.Draw(canvas)
    d.text((margin, 30), "整车多角度实物照片", fill=(22, 44, 72), font=FONT_TITLE)
    d.text((margin, 84), "展示车模底盘、海鸥派 SS928、STM32 控制板、摄像头、供电与线束的整体集成状态", fill=(85, 99, 118), font=FONT_NOTE)
    y1 = margin + title_h
    for idx in range(3):
        draw_cell(canvas, (margin + idx * (top_w + gap), y1), (top_w, top_h), PHOTOS[idx][0], PHOTOS[idx][1])
    y2 = y1 + top_h + gap
    for j, idx in enumerate([3, 4]):
        draw_cell(canvas, (margin + j * (bot_w + gap), y2), (bot_w, bot_h), PHOTOS[idx][0], PHOTOS[idx][1])
    d.text((margin, H - foot_h + 18), "注：照片均为实物拍摄，用于支撑机械成果、硬件集成和现场展示材料。", fill=(85, 99, 118), font=FONT_NOTE)
    # Save as JPEG to keep DOCX size moderate and strip EXIF metadata.
    canvas.save(COLLAGE, quality=92, optimize=True)


def set_run_font(run, name="微软雅黑", size=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_text(paragraph, text, size=None, bold=None, italic=None):
    for r in paragraph.runs:
        r.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    set_run_font(run, size=size, bold=bold, italic=italic)


def move_after(new_para, cursor_para):
    cursor_para._p.addnext(new_para._p)
    return new_para


def add_picture_block_after(doc, after_para, image_path):
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(6)
    intro.paragraph_format.space_after = Pt(4)
    intro.paragraph_format.keep_with_next = True
    r = intro.add_run("图证补充：整车照片从正面、俯视、侧面和后部展示了车体底盘、板卡分层安装、摄像头固定、供电与线束布置，可作为工程成果与现场展示的实物依据。")
    set_run_font(r, size=10.5)
    cursor = move_after(intro, after_para)

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(3)
    pic_p.paragraph_format.space_after = Pt(2)
    pic_p.paragraph_format.keep_with_next = True
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Inches(6.25))
    # Accessibility/alt text for the inserted drawing.
    for drawing in run._element.xpath(".//w:drawing"):
        doc_pr = drawing.xpath(".//wp:docPr")
        if doc_pr:
            doc_pr[0].set("name", "整车多角度实物照片")
            doc_pr[0].set("descr", "正面、左前斜上、俯视、侧视和后部五张实物照片组成的图版，展示自动泊车车模的硬件集成状态。")
    cursor = move_after(pic_p, cursor)

    cap = doc.add_paragraph()
    try:
        cap.style = doc.styles["Caption"]
    except Exception:
        cap.style = doc.styles["Normal"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(9)
    set_paragraph_text(cap, "图2 整车多角度实物照片：展示车模底盘、海鸥派 SS928、STM32 控制板、摄像头、供电与线束等硬件集成状态。", size=9, italic=True)
    move_after(cap, cursor)


make_collage()

doc = Document(str(SRC))

# Strengthen the engineering outcome table evidence cell.
for table in doc.tables:
    for row in table.rows:
        if row.cells and "机械与底盘" in row.cells[0].text:
            target = row.cells[2].paragraphs[0]
            set_paragraph_text(target, "整车多角度实物照片、底盘实物、标定数据表、曲率表。", size=9)

insert_after = None
for p in doc.paragraphs:
    if "表8 工程成果汇总" in p.text:
        insert_after = p
        break
if insert_after is None:
    # Fallback: place after the 3.2 heading paragraph if the caption is absent.
    for p in doc.paragraphs:
        if p.text.strip() == "3.2 工程成果":
            insert_after = p
            break
if insert_after is None:
    raise RuntimeError("未找到照片插入位置：表8 工程成果汇总 / 3.2 工程成果")

add_picture_block_after(doc, insert_after, COLLAGE)

# Keep metadata neutral.
cp = doc.core_properties
cp.title = "基于海鸥派 SS928 与 STM32 底盘的视觉闭环自主泊车系统"
cp.subject = "竞赛作品报告"
cp.author = ""
cp.comments = ""
cp.keywords = ""

doc.save(str(OUT))
print(f"OUT={OUT}")
print(f"COLLAGE={COLLAGE}")
