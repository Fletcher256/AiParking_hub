from docx import Document
from pathlib import Path
p=Path(r'C:\Users\Cheng\Desktop\嵌入式大赛作品报告_视觉闭环自主泊车系统_增强版_含流程图.docx')
doc=Document(str(p))
print('PARAGRAPHS')
for i,p in enumerate(doc.paragraphs):
    t=' '.join(p.text.split())
    if t:
        print(f'P{i:04d} [{p.style.name}] {t}')
print('\nTABLES')
for ti,tbl in enumerate(doc.tables):
    print(f'-- TABLE {ti+1} --')
    for ri,row in enumerate(tbl.rows):
        cells=[' '.join(c.text.split()) for c in row.cells]
        print(f'R{ri}: ' + ' | '.join(cells))
