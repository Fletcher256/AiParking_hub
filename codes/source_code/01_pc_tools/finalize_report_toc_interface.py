from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

SRC = Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版_最终补图表版.docx")
OUT = Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版_最终提交整理版.docx")
INTERFACE_IMG = Path(r"D:/parking_board_agent/scratch_doc_review/stm32_pdf_render/stm32_controller.png")

if not SRC.exists():
    raise FileNotFoundError(SRC)
if not INTERFACE_IMG.exists():
    raise FileNotFoundError(INTERFACE_IMG)


def set_run_font(run, name="微软雅黑", size=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_text(paragraph, text, size=None, bold=None, italic=None):
    for r in paragraph.runs:
        r.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    set_run_font(run, size=size, bold=bold, italic=italic)


def move_after(new_para, cursor_para):
    cursor_para._p.addnext(new_para._p)
    return new_para


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def add_interface_block(doc, after_para):
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(6)
    intro.paragraph_format.space_after = Pt(4)
    intro.paragraph_format.keep_with_next = True
    r = intro.add_run("图证补充：STM32 底盘控制接口连接图给出了电机 PWM、编码器、舵机 PWM、BIM270、OLED、蓝牙调试模块、指示灯以及 SS928 USB 串口通信等关键输入输出信号。")
    set_run_font(r, size=10.5)
    cursor = move_after(intro, after_para)

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(3)
    pic_p.paragraph_format.space_after = Pt(2)
    pic_p.paragraph_format.keep_with_next = True
    run = pic_p.add_run()
    run.add_picture(str(INTERFACE_IMG), width=Inches(6.25))
    for drawing in run._element.xpath(".//w:drawing"):
        doc_pr = drawing.xpath(".//wp:docPr")
        if doc_pr:
            doc_pr[0].set("name", "STM32 底盘控制接口连接图")
            doc_pr[0].set("descr", "STM32C8T6 控制板接口连接图，标注电机 PWM、编码器、舵机 PWM、BIM270、OLED、蓝牙调试模块、指示灯和 SS928 USB 串口通信信号。")
    cursor = move_after(pic_p, cursor)

    cap = doc.add_paragraph()
    try:
        cap.style = doc.styles["Caption"]
    except Exception:
        cap.style = doc.styles["Normal"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(9)
    set_paragraph_text(cap, "图2 STM32 底盘控制接口连接图：标注电机 PWM、编码器、舵机 PWM、BIM270、OLED、蓝牙调试模块、指示灯和 SS928 USB 串口通信信号。", size=9, italic=True)
    move_after(cap, cursor)


def normalize_caption(paragraph, old_prefix, new_prefix):
    txt = paragraph.text
    if txt.startswith(old_prefix):
        set_paragraph_text(paragraph, txt.replace(old_prefix, new_prefix, 1), size=9, italic=True)


doc = Document(str(SRC))

# Replace manual TOC entries with a Word TOC placeholder. Keep the visible "目录" title but do not let it be included as Heading 1.
toc_idx = None
first_part_idx = None
for idx, p in enumerate(doc.paragraphs):
    if p.text.strip() == "目录":
        toc_idx = idx
    if toc_idx is not None and p.text.strip().startswith("第一部分") and p.style.name.startswith("Heading"):
        first_part_idx = idx
        break
if toc_idx is not None and first_part_idx is not None:
    toc_p = doc.paragraphs[toc_idx]
    try:
        toc_p.style = doc.styles["Normal"]
    except Exception:
        pass
    toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(toc_p, "目录", size=16, bold=True)
    # Delete old manual TOC entries between title and first body heading.
    for p in list(doc.paragraphs[toc_idx + 1:first_part_idx]):
        delete_paragraph(p)
    placeholder = doc.add_paragraph("__AUTO_TOC_PLACEHOLDER__")
    placeholder.paragraph_format.space_after = Pt(8)
    move_after(placeholder, toc_p)
else:
    raise RuntimeError("未找到目录或第一部分标题，无法自动目录占位")

# Insert interface figure after the paragraph under 2.2.3.
insert_after = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "2.2.3 电路与接口说明":
        for q in doc.paragraphs[i+1:]:
            if q.text.strip() and not q.style.name.startswith("Heading"):
                insert_after = q
                break
        break
if insert_after is None:
    raise RuntimeError("未找到 2.2.3 电路与接口说明后的插入位置")
add_interface_block(doc, insert_after)

# Renumber figure captions into document order after adding interface figure.
for p in doc.paragraphs:
    normalize_caption(p, "图2 系统整体架构图", "图1 系统整体架构图")
    normalize_caption(p, "图1 line_follow", "图3 line_follow")
    normalize_caption(p, "图3 整车多角度实物照片", "图4 整车多角度实物照片")
    normalize_caption(p, "图4 泊车前后对比", "图5 泊车前后对比")
    normalize_caption(p, "图5 YOLO 车位识别效果", "图6 YOLO 车位识别效果")

# Update references in paragraphs and tables.
repls = [
    ("软件系统章节的图1用于说明 line_follow 决策逻辑", "软件系统章节的图3用于说明 line_follow 决策逻辑"),
    ("泊车前后对比见图4", "泊车前后对比见图5"),
    ("可结合图4泊车前后对比复核", "可结合图5泊车前后对比复核"),
    ("识别效果见图5", "识别效果见图6"),
]
for p in doc.paragraphs:
    txt = p.text
    new = txt
    for a, b in repls:
        new = new.replace(a, b)
    if new != txt:
        set_paragraph_text(p, new)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                txt = p.text
                new = txt
                for a, b in repls:
                    new = new.replace(a, b)
                if new != txt:
                    set_paragraph_text(p, new, size=9)

# Strengthen evidence entries for interface diagram.
for table in doc.tables:
    for row in table.rows:
        if row.cells and "电路与通信" in row.cells[0].text and len(row.cells) >= 3:
            set_paragraph_text(row.cells[2].paragraphs[0], "STM32 底盘控制接口连接图、串口日志、STM32 状态字段、执行反馈。", size=9)

cp = doc.core_properties
cp.title = "基于海鸥派 SS928 与 STM32 底盘的视觉闭环自主泊车系统"
cp.subject = "竞赛作品报告"
cp.author = ""
cp.comments = ""
cp.keywords = ""

doc.save(str(OUT))
print(f"OUT={OUT}")
