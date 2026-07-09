from docx import Document
from pathlib import Path
p=Path(r'C:\Users\Cheng\Desktop\嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版.docx')
doc=Document(str(p))
for i,p in enumerate(doc.paragraphs):
    t=' '.join(p.text.split())
    if t:
        print(f'P{i:04d} [{p.style.name}] {t}')
print('\nBAD SEARCH')
words=['建议','正式提交','本文档','不插入','自报','盲','答辩','心得','国赛','调权重','调参数','fallback','控制日志自报','补充计划','controller','拉回电脑','console','unittest']
alltexts=[]
for p in doc.paragraphs: alltexts.append(p.text)
for ti,tbl in enumerate(doc.tables):
    for row in tbl.rows:
        for cell in row.cells:
            alltexts.append(cell.text)
for w in words:
    hits=[t for t in alltexts if w in t]
    if hits:
        print('WORD', w)
        for h in hits[:10]: print('  ', ' '.join(h.split()))
print('\nTABLES')
for ti,tbl in enumerate(doc.tables):
    print(f'-- TABLE {ti+1} --')
    for ri,row in enumerate(tbl.rows):
        print(f'R{ri}: ' + ' | '.join(' '.join(c.text.split()) for c in row.cells))
