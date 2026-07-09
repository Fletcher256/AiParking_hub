from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = Path(r"C:\Users\Cheng\Desktop\嵌入式大赛作品报告_视觉闭环自主泊车系统_增强版.docx")
IMG = Path(r"D:\parking_board_agent\docs\line_follow_flowchart.png")
OUT = Path(r"C:\Users\Cheng\Desktop\嵌入式大赛作品报告_视觉闭环自主泊车系统_增强版_含流程图.docx")

if not SRC.exists():
    raise FileNotFoundError(SRC)
if not IMG.exists():
    raise FileNotFoundError(IMG)

doc = Document(str(SRC))

# Find the paragraph immediately after “2.3.1 软件整体介绍”.
insert_after = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "2.3.1 软件整体介绍":
        for q in doc.paragraphs[i+1:]:
            if q.text.strip():
                insert_after = q
                break
        break
if insert_after is None:
    raise RuntimeError("未找到 2.3.1 软件整体介绍后的插入位置")

# Do not duplicate insertion if script is rerun on an already-edited copy.
for p in doc.paragraphs:
    if "图1 line_follow 倒车决策核心流程图" in p.text:
        raise RuntimeError("文档中已存在该流程图说明，避免重复插入")


def set_run_font(run, name="微软雅黑", size=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def move_after(new_para, cursor):
    cursor._p.addnext(new_para._p)
    return new_para

cursor = insert_after

pic_p = doc.add_paragraph()
pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic_p.paragraph_format.space_before = Pt(4)
pic_p.paragraph_format.space_after = Pt(3)
pic_p.paragraph_format.keep_with_next = True
pic_p.paragraph_format.keep_together = True
pic_run = pic_p.add_run()
pic_run.add_picture(str(IMG), width=Inches(5.0))
# Add accessible title/alt text to the inserted drawing.
for drawing in pic_run._element.xpath('.//w:drawing'):
    doc_pr = drawing.xpath('.//wp:docPr')
    if doc_pr:
        doc_pr[0].set('name', 'line_follow 倒车决策核心流程图')
        doc_pr[0].set('descr', 'line_follow 倒车决策核心流程：闭环预演纯倒车可达性，不可达时搜索前进 shuffle 后重新规划。')
cursor = move_after(pic_p, cursor)

cap = doc.add_paragraph()
try:
    cap.style = doc.styles["Caption"]
except Exception:
    cap.style = doc.styles["Normal"]
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_before = Pt(0)
cap.paragraph_format.space_after = Pt(9)
cap.paragraph_format.keep_together = True
cap_run = cap.add_run(
    "图1 line_follow 倒车决策核心流程图：每步动作先闭环 rollout 预演纯倒车可达性；"
    "不可达时搜索前进 shuffle，再回到观察—决策—执行循环。"
)
set_run_font(cap_run, size=9, italic=True)
cursor = move_after(cap, cursor)

# Remove contradictions caused by adding a design-flow figure.
for p in doc.paragraphs:
    txt = p.text.strip()
    if "本文档不插入自动生成图片，避免图像质量参差" in txt:
        for r in p.runs:
            r.text = ""
        r = p.add_run(
            "为增强国赛答辩中的证据说服力，正式提交前建议补充 3 组不同初始姿态的重复实验，"
            "并同步记录外部尺量结果。本文档已在软件系统章节插入 line_follow 决策流程图用于说明控制逻辑；"
            "实车图像、视频和截图可在最终排版阶段从真实实验材料中选取。"
        )
        set_run_font(r, size=10.5)
    elif "本文档不插入自动生成图片；需要图像时应使用真实实验截图" in txt:
        for r in p.runs:
            r.text = ""
        r = p.add_run(
            "以下材料建议作为最终提交或答辩附件准备。除软件系统章节的决策流程图外，"
            "图证材料建议优先使用真实实验截图、实车照片或视频关键帧。"
        )
        set_run_font(r, size=10.5)

doc.save(str(OUT))
print(OUT)
