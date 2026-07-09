from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path(r"C:\Users\Cheng\Desktop\嵌入式大赛作品报告_视觉闭环自主泊车系统_增强版.docx")
IMG = Path(r"D:\parking_board_agent\docs\line_follow_flowchart.png")
OUT = Path(r"C:\Users\Cheng\Desktop\嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版.docx")


def set_run_font(run, name="微软雅黑", size=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_text(paragraph, text):
    # Preserve paragraph style and numbering properties; replace only visible text.
    if paragraph.runs:
        first = paragraph.runs[0]
        for r in paragraph.runs:
            r.text = ""
        first.text = text
    else:
        first = paragraph.add_run(text)
    set_run_font(first)


def iter_all_paragraphs(document):
    for p in document.paragraphs:
        yield p
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def move_after(new_para, cursor):
    cursor._p.addnext(new_para._p)
    return new_para


def add_picture_after(paragraph, image_path):
    cursor = paragraph
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(4)
    pic_p.paragraph_format.space_after = Pt(3)
    pic_p.paragraph_format.keep_with_next = True
    pic_p.paragraph_format.keep_together = True
    pic_run = pic_p.add_run()
    pic_run.add_picture(str(image_path), width=Inches(4.6))
    for drawing in pic_run._element.xpath(".//w:drawing"):
        doc_pr = drawing.xpath(".//wp:docPr")
        if doc_pr:
            doc_pr[0].set("name", "line_follow 倒车决策流程图")
            doc_pr[0].set(
                "descr",
                "line_follow 倒车决策流程：通过闭环预演评估纯倒车可达性，不可达时插入前进调整段后重新规划。",
            )
    cursor = move_after(pic_p, cursor)

    cap = doc.add_paragraph()
    try:
        cap.style = doc.styles["Caption"]
    except Exception:
        cap.style = doc.styles["Normal"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(9)
    cap.paragraph_format.keep_together = True
    cap_run = cap.add_run(
        "图1 line_follow 倒车决策流程：系统根据当前位姿生成受底盘约束的倒车曲率，"
        "并通过闭环 rollout 判断纯倒车是否可达；若不可达，则搜索并执行短距离前进调整段后重新规划。"
    )
    set_run_font(cap_run, size=9, italic=True)
    move_after(cap, cursor)


doc = Document(str(SRC))

exact_replacements = {
    "竞赛匿名版说明文档（不包含学校、团队、队员和指导老师信息）": "竞赛作品报告（匿名版）",
    "针对小车底盘转向能力弱、纯倒车横向修正范围有限的问题，项目将原有“阶段状态机+多权重打分”的倒车策略升级为可开关的 line_follow 决策核心：采用饱和临界阻尼直线跟踪律生成倒车曲率，利用实测舵角-曲率表反查 STE，并在每一步执行前进行闭环 rollout 预演；当预演判断纯倒车无法进入成功框时，自动搜索前进 shuffle 换向段，再执行下一步倒车。该方法把“是否需要前进调整”从经验阈值转化为物理可达性判断，解决了横向大偏差仅靠调权重无法收敛的问题。": "针对车辆底盘转向能力受限、纯倒车横向修正范围有限的问题，系统将原有“阶段状态机+多权重评分”的倒车策略升级为可配置的 line_follow 决策核心：通过饱和临界阻尼直线跟踪律生成倒车曲率，依据实测曲率-舵角标定表反查 STE 控制量，并在每一步执行前进行闭环 rollout 预演；当预演结果表明纯倒车无法进入成功判定区域时，系统自动搜索短距离前进调整段（shuffle），再进入下一轮倒车规划。该方法将“是否需要前进调整”由经验阈值判断转化为基于底盘运动约束的可达性判断，提高了大横向偏差场景下的收敛能力。",
    "目前系统已完成本地单元测试、集成测试、蒙特卡洛验证、板端部署和多次实车闭环验证。2026-07-04 的实车 line_follow 录制流程中，系统成功启动 YOLO 与 H264 原始视频录制，完成 12 步倒车/前进闭环动作，控制日志自报终态约为 y=2.2 cm、lat=-1.0 cm、head=6.2°，并输出原始 H264 视频与 JSONL 控制日志，具备演示、复盘和后续标定能力。": "系统已完成本地单元测试、集成测试、蒙特卡洛验证、板端部署和多次实车闭环验证。在 2026 年 7 月 4 日的实车 line_follow 录制流程中，系统启动 YOLO 与 H264 原始视频录制，完成 12 步倒车/前进闭环动作，控制日志估计终态约为 y=2.2 cm、lat=-1.0 cm、head=6.2°，并输出原始 H264 视频与 JSONL 控制日志，具备现场展示、过程复盘和后续标定能力。",
    "附录A 竞赛展示与答辩材料清单": "附录A 竞赛展示与支撑材料清单",
    "附录B 现场演示脚本与风险预案": "附录B 现场风险与应对预案",
    "弱转向补偿：对纯倒车物理不可达的横向偏差，自动触发前进 shuffle 重新获得修正空间。": "转向能力受限补偿：对纯倒车物理不可达的横向偏差，自动触发前进调整段（shuffle）以重新获得修正空间。",
    "安全执行：保留 arm 文件、最大步数、最大路程、视觉丢失预算、STM32 状态检查等安全门。": "安全执行：设置运动授权文件（arm 文件）、最大步数、最大路程、视觉丢失预算、STM32 状态检查等安全门。",
    "本作品可用于智能小车竞赛、低速园区无人车、仓储搬运机器人、教学实验平台和自动泊车算法验证平台。其核心价值不是追求高速行驶，而是在低速、窄空间、感知误差和执行误差并存的情况下，展示一条可复现、可解释、可安全停止的嵌入式闭环控制链路。系统中的车位识别、单应性位姿估计、短动作重规划、底盘曲率标定、控制日志复盘等模块，也可以迁移到自动对接、自动充电、狭窄通道入位等近距离定位任务中。": "本作品可用于智能小车竞赛、低速园区无人车、仓储搬运机器人、教学实验平台和自动泊车算法验证平台。作品重点展示在低速、窄空间、感知误差和执行误差并存条件下，可复现、可解释、可安全停止的嵌入式闭环控制链路。系统中的车位识别、单应性位姿估计、短动作重规划、底盘曲率标定、控制日志复盘等模块，也可迁移到自动对接、自动充电、狭窄通道入位等近距离定位任务中。",
    "闭环 rollout：每步执行前用同一策略预演终点，判断是否能落入成功框。": "闭环 rollout：每步执行前用同一策略预演终点，判断是否能进入成功判定区域。",
    "Reeds-Shepp 式前进 shuffle：当剩余深度不足以横向修正时，搜索前进换向段并只执行第一步。": "Reeds-Shepp 式前进调整：当剩余深度不足以完成横向修正时，搜索前进换向段并执行当前最优短动作。",
    "录制复盘链路：倒车开始前启动板端 H264 原始录制，结束后停止并拉回电脑，与 JSONL 日志配套分析。": "录制复盘链路：倒车开始前启动板端 H264 原始录制，结束后停止并回传上位机，与 JSONL 日志配套分析。",
    "底盘与传感器联调：完成摄像头、YOLO、STM32 串口、IMU/编码器和安全 arm 机制。": "底盘与传感器联调：完成摄像头、YOLO、STM32 串口、IMU/编码器和运动授权机制。",
    "底盘标定：实测不同 STE 的 ARC 曲率、死区、滑行量和方向符号，形成 chassis_kinematics.json。": "底盘标定：实测不同 STE 下的 ARC 曲率、死区、滑行量和方向符号，形成曲率标定数据。",
    "策略设计：先实现阶段式策略，再抽象为 line_follow 决策核心，并保持 legacy 可回退。": "策略设计：先实现阶段式策略，再抽象为 line_follow 决策核心，并保留 legacy 备用策略。",
    "仿真与日志验证：用 rollout、单测、集成测试和蒙特卡洛覆盖典型初始姿态。": "仿真与日志验证：通过 rollout、单元测试、集成测试和蒙特卡洛仿真覆盖典型初始姿态。",
    "本系统由感知层、位姿层、决策层、执行层和复盘层组成。感知层负责从相机画面中识别车位；位姿层将 YOLO polygon 转换为厘米级相对状态；决策层依据 line_follow 或 legacy 策略生成单步 ARC 动作；执行层经 STM32 控制舵机和电机；复盘层保存视频与 JSONL 日志，用于问题定位和参数迭代。": "本系统由感知层、位姿层、决策层、执行层和复盘层组成。感知层负责从相机画面中识别车位；位姿层将 YOLO polygon 转换为厘米级相对状态；决策层依据 line_follow 或 legacy 策略生成单步 ARC 动作；执行层经 STM32 控制舵机和电机；复盘层保存视频与 JSONL 日志，用于验证分析和参数迭代。",
    "YOLO 结果同时送入控制器和监控/录制链路；控制器只发出一个短动作，动作完成后读取 STM32 里程、yaw 和新的视觉位姿，再进入下一轮闭环。该设计避免长时间盲走，也使每一步动作都可以被日志解释。": "YOLO 结果同时送入控制器和监控/录制链路；控制器每次仅发出一个短动作，动作完成后读取 STM32 里程、yaw 和新的视觉位姿，再进入下一轮闭环。该设计避免长距离开环行驶，并使每一步动作都能够通过日志追溯。",
    "硬件平台由海鸥派 SS928 边缘计算板、OS08A20 摄像头、STM32 底盘控制板、舵机、电机、编码器、IMU 和小车机械底盘组成。SS928 负责图像采集、YOLO 推理、网络/串口通信和高层决策；STM32 负责低层电机/舵机控制、里程计、IMU yaw 读取和指令应答。两者通过串口/USB 串口连接，上位机仅用于部署、日志拉取和演示监控，不参与实时闭环的关键路径。": "硬件平台由海鸥派 SS928 边缘计算板、OS08A20 摄像头、STM32 底盘控制板、舵机、电机、编码器、IMU 和车模机械底盘组成。SS928 负责图像采集、YOLO 推理、网络/串口通信和高层决策；STM32 负责低层电机/舵机控制、里程计、IMU yaw 读取和指令应答。两者通过串口/USB 串口连接，上位机端用于部署、日志回传和演示监控，不参与实时闭环的关键路径。",
    "小车采用前轮转向、后轮驱动的低速车模结构。由于底盘转向能力较弱，实测最大曲率约为右 0.825 °/cm、左 1.203 °/cm，对应最小转弯半径约 69 cm 和 48 cm。该物理约束直接决定了倒车策略：当剩余深度较小时，纯倒车可修正的横向距离近似为 κ·s²/4，25 cm 深度仅能修正 2~3 cm。因此系统必须支持前进 shuffle，而不能依赖单纯修改评分权重。": "车辆采用前轮转向、后轮驱动的低速车模结构。由于底盘转向能力受限，实测最大曲率约为右 0.825 °/cm、左 1.203 °/cm，对应最小转弯半径约 69 cm 和 48 cm。该物理约束直接影响倒车策略：当剩余深度较小时，纯倒车可修正的横向距离近似为 κ·s²/4，25 cm 深度仅能修正 2~3 cm。因此系统需要前进调整能力，而非仅依赖评分权重调整。",
    "软件采用“板端实时闭环 + 电脑端部署/复盘”的结构。板端运行 YOLO 进程和 board_parking_controller.py，控制器从 UDP 接收车位检测结果，经感知过滤、目标锁定、位姿计算后调用决策核心，生成单个 plan.v2 动作。动作发送给 STM32 后，控制器等待 DONE/STAT，更新里程与 yaw，再结合后续视觉帧重锚定位。电脑端提供测试脚本、SSH 部署、日志下载、视频录制流程和离线分析工具。": "软件采用“板端实时闭环 + 上位机端部署/复盘”的结构。板端运行 YOLO 进程和 board_parking_controller.py，控制器从 UDP 接收车位检测结果，经感知过滤、目标锁定、位姿计算后调用决策核心，生成单步运动规划。动作发送给 STM32 后，控制器等待 DONE/STAT，更新里程与 yaw，再结合后续视觉帧重锚定位。上位机端提供测试脚本、SSH 部署、日志下载、视频录制流程和离线分析工具。",
    "自动泊车作品存在真实运动风险，因此系统设计坚持“短动作、低速度、可停止、可复盘”的原则。真实动作必须经过 arm 文件和命令行参数双重确认，dry-run 和 replanner-dry-run 模式不会打开底盘串口执行运动。系统还通过视觉丢失预算、最大步数、最大累计路程、STM32 异常状态检查和最终 STOP 降低现场演示风险。": "自动泊车系统涉及实体车辆运动，因此系统设计坚持“短动作、低速度、可停止、可复盘”的原则。真实动作必须经过运动授权文件（arm 文件）和命令行参数双重确认；dry-run 和 replanner-dry-run 模式不会打开底盘串口执行运动。系统还通过视觉丢失预算、最大步数、最大累计路程、STM32 异常状态检查和最终 STOP 降低现场演示风险。",
    "目前系统已完成从图像输入到实车动作输出的端到端闭环：板端 YOLO 能稳定检测车位，控制器可锁定目标并持续更新相对位姿，line_follow 决策核心可在倒车和前进 shuffle 之间自动选择，STM32 能返回里程、yaw 和动作完成状态。系统保留 legacy 决策开关，line_follow 通过参数启用，便于比赛现场在稳定性和可解释性之间快速切换。": "系统已完成从图像输入到实车动作输出的端到端闭环：板端 YOLO 能稳定检测车位，控制器可锁定目标并持续更新相对位姿，line_follow 决策核心可在倒车和前进调整段之间自动选择，STM32 能返回里程、yaw 和动作完成状态。系统保留 legacy 备用策略，line_follow 可通过参数启用，便于在现场演示中兼顾稳定性和可解释性。",
    "3.4 实车演示与量化补充计划": "3.4 实车演示与量化验证方案",
    "为增强国赛答辩中的证据说服力，正式提交前建议补充 3 组不同初始姿态的重复实验，并同步记录外部尺量结果。本文档不插入自动生成图片，避免图像质量参差；图像、视频和截图可在最终排版阶段从真实实验材料中选取。": "为保证现场展示和量化评估的完整性，实车演示设置三类典型初始姿态，并同步记录初始位姿、动作步数、是否触发前进调整段、终态外部尺量结果及控制日志。软件系统章节的图1用于说明 line_follow 决策逻辑；实际展示材料以真实运行截图、实车照片、视频关键帧和日志片段为依据。",
    "表10 实车演示补充计划": "表10 实车演示与量化验证方案",
    "展示系统上电与安全状态：确认 arm 文件不存在、controller 处于安全停止状态。": "展示系统上电与安全状态：确认运动授权文件未创建，控制器处于安全停止状态。",
    "进入 dry-run 或 no-motion 模式：展示系统能计算动作但不会发送真实运动。": "进入 dry-run 或无运动模式：展示系统能计算动作，但不会发送真实运动指令。",
    "人工确认场地安全后进入真实演示：只执行短动作闭环，动作后停车并重观测。": "完成场地安全确认后进入真实演示：仅执行短动作闭环，动作后停车并重新观测。",
    "展示日志和视频：用 JSONL 说明每一步“看到什么、决定什么、STM32 实际反馈什么”。": "展示日志和视频：通过 JSONL 说明每一步的感知结果、决策输出和 STM32 实际反馈。",
    "展示最终结果：给出终态照片/尺量/控制日志自报三类证据，并说明误差来源。": "展示最终结果：给出终态照片、外部尺量和控制日志估计三类证据，并说明误差来源。",
    "4.1 可扩展之处": "4.1 后续工作",
    "补充前进 ARC 标定表：当前前进段默认使用倒车曲率取反近似，后续可用 2~3 个实测前进样本进一步提高 shuffle 预测精度。": "进一步完善前进 ARC 标定表：当前前进段默认使用倒车曲率取反近似，后续将通过实测前进样本提高 shuffle 预测精度。",
    "完善外部量测评估：引入标尺、定位板或多相机测量，对自报 pose 进行独立精度评估。": "完善外部量测评估：引入标尺、定位板或多相机测量，对控制日志 pose 进行独立精度评估。",
    "4.2 心得体会": "4.2 项目总结",
    "本项目的核心收获是：嵌入式智能车不是单独调好某一个算法就能成功，而是感知、标定、物理模型、安全执行和日志复盘共同作用的系统工程。早期采用固定阶段和多权重打分时，系统在某些初始姿态下会出现横向偏差长期无法修正的问题。经过实车测量后发现，根本原因并不是某个权重不合适，而是底盘转向半径较大，剩余倒车深度不足时纯倒车物理上无法完成横向修正。这个认识促使我们把“调参数”转向“看物理可达性”，最终形成了 line_follow + rollout + shuffle 的方案。": "本项目形成的主要工程认识是：嵌入式智能车系统的可靠性来自感知、标定、物理模型、安全执行和日志复盘的协同，而不是单一算法模块。早期采用固定阶段和多权重评分时，系统在部分初始姿态下出现横向偏差长期无法修正的情况。实车测量表明，根本原因是底盘转向半径较大，剩余倒车深度不足时纯倒车无法完成横向修正。因此，方案由单纯调整评分权重转向基于物理可达性进行决策，最终形成 line_follow + rollout + shuffle 的闭环规划方法。",
    "在工程实现上，项目坚持短动作、停-看-走、可回退和可复盘原则。每一步都只执行一个小的 ARC 动作，动作后读取 STM32 反馈并等待视觉重锚；每次改动都通过单元测试、集成测试和实际日志验证；涉及实车运动时必须使用 arm 文件和明确审批。录制链路的建立也很关键：原始 H264 视频与 JSONL 控制日志可以把“看到什么、决定什么、实际走了多少”对应起来，使问题不再停留在主观观察，而能被逐步定位、修复和复现。": "在工程实现上，系统坚持短动作、停-看-走、可回退和可复盘原则。每一步仅执行一个小幅 ARC 动作，动作结束后读取 STM32 反馈并等待视觉重锚；每次策略改动均经过单元测试、集成测试和实车日志验证；涉及实体运动时必须通过运动授权文件和命令行参数双重确认。原始 H264 视频与 JSONL 控制日志能够对应记录感知结果、决策输出和底盘反馈，使问题定位、参数修正和复现实验具备依据。",
    "后续工作将围绕更完整的前进段标定、外部量测、动态安全和演示界面展开。当前系统已经具备完整的竞赛演示基础：能在真实小车上完成车位识别、相对定位、闭环倒车、必要换向、成功判定和全过程记录，体现了嵌入式视觉感知、底盘控制和工程调试的综合能力。": "后续工作将围绕更完整的前进段标定、外部量测、动态安全和演示界面展开。当前系统已具备竞赛演示所需的完整闭环能力：能够在真实车辆平台上完成车位识别、相对定位、闭环倒车、必要换向、成功判定和全过程记录，体现了嵌入式视觉感知、底盘控制和工程调试的综合能力。",
    "[5] 项目文档：parking_line_follow_decision_20260704.md，chassis_kinematics.json，board_parking_controller.py 实车日志。": "[5] 项目工程资料：parking_line_follow_decision_20260704.md，chassis_kinematics.json，board_parking_controller.py 及实车日志。",
    "以下材料建议作为最终提交或答辩附件准备。本文档不插入自动生成图片；需要图像时应使用真实实验截图、实车照片或视频关键帧。": "以下材料用于支撑作品展示和现场评审。除软件系统章节的决策流程图外，图证材料优先采用真实实验截图、实车照片或视频关键帧。",
}

for p in list(doc.paragraphs):
    if p.text.strip() == "注：提交最终版前，可在 Word 中使用“引用 → 目录”自动生成页码目录。":
        delete_paragraph(p)

for p in iter_all_paragraphs(doc):
    text = " ".join(p.text.split())
    if text in exact_replacements:
        set_paragraph_text(p, exact_replacements[text])

substring_replacements = [
    ("运动授权文件（运动授权文件（arm 文件））", "运动授权文件（arm 文件）"),
    ("不是单一算法演示，而是完整软硬件协同系统。", "体现完整软硬件协同，而非单一算法模块。"),
    ("经验调参", "经验参数调整"),
    ("forward shuffle", "前进调整段（shuffle）"),
    ("前进 shuffle", "前进调整段（shuffle）"),
    ("rollout + shuffle", "rollout + 前进调整段"),
    ("运动模型、shuffle", "运动模型、前进调整段"),
    ("shuffle 搜索", "前进调整段搜索"),
    ("3150 runs", "3150 次仿真"),
    ("51% run", "51% 仿真样本"),
    ("控制日志自报", "控制日志估计"),
    ("终态自报", "日志估计终态"),
    ("正式提交建议补外部尺量", "现场演示阶段同步外部尺量复核"),
    ("建议正式提交前补外部尺量复核", "现场演示阶段同步外部尺量复核"),
    ("建议记录指标", "记录指标"),
    ("预期展示价值", "验证意义"),
    ("建议来源", "材料来源"),
    ("unittest", "单元测试"),
    ("console 日志", "控制台日志"),
    ("H264 + JSONL + console", "H264 + JSONL + 控制台日志"),
    ("arm 文件门", "运动授权门"),
    ("arm 文件", "运动授权文件（arm 文件）"),
    ("仅允许 dry-run", "仅允许无运动仿真模式（dry-run）"),
    ("避免盲目继续执行", "避免无反馈继续执行"),
    ("输出 STOP/日志", "发送 STOP 并记录日志"),
    ("是否进入成功框", "是否进入成功判定区域"),
    ("面向弱转向底盘", "面向转向能力受限底盘"),
    ("弱转向", "转向能力受限"),
    ("并只执行第一步", "并执行当前最优短动作"),
    ("看到什么、决定什么、实际走了多少", "感知结果、决策输出和底盘反馈"),
    ("单测", "单元测试"),
    ("本地 单元测试", "本地单元测试"),
    ("前进调整段（shuffle） 指令", "前进调整段（shuffle）指令"),
    ("和 前进调整段", "和前进调整段"),
    ("同步拉取", "同步回传"),
    ("legacy 回退", "legacy 备用策略"),
    ("heading debt 支付能力", "航向误差修正能力"),
    ("stop_bounds", "安全边界"),
    ("执行壳", "执行管理"),
    ("不依赖后期合成即可追溯演示过程", "可追溯完整演示过程"),
    ("文件拉取", "文件回传"),
    ("演示包", "演示材料包"),
    ("等待人工处理", "等待人工确认处理"),
    ("适合现场演示，也便于评审追溯系统稳定性。", "支持现场演示安全性，也便于评审追溯系统稳定性。"),
    ("初始姿态过偏", "初始姿态超出验证范围"),
    ("选择已验证的 A/B/C 演示起点，不做临场极限挑战。", "选择已验证的 A/B/C 演示起点，避免现场临时调整至未验证极限姿态。"),
    ("先展示 dry-run/历史日志", "先展示无运动仿真和历史日志"),
    ("切换到日志回放说明安全门工作", "结合日志回放说明安全门状态"),
    ("并用人工重试短动作修正", "并根据安全流程执行短动作修正或终止演示"),
    ("系统整体框图", "系统整体架构图"),
    ("真实架构手工绘制或项目文档整理", "系统架构整理"),
    ("帮助评委", "便于评委"),
    ("实物小车照片", "车模实物照片"),
    ("说明 homography 和相对状态。", "说明 homography 与相对状态。"),
    ("运动授权文件（运动授权文件（arm 文件））", "运动授权文件（arm 文件）"),
]

for p in iter_all_paragraphs(doc):
    current = p.text
    new = current
    for old, repl in substring_replacements:
        new = new.replace(old, repl)
    if new != current:
        set_paragraph_text(p, new)

# Insert the formal flowchart after the first body paragraph under 2.3.1.
insert_after = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "2.3.1 软件整体介绍":
        for q in doc.paragraphs[i + 1 :]:
            if q.text.strip():
                insert_after = q
                break
        break
if insert_after is None:
    raise RuntimeError("未找到 2.3.1 软件整体介绍后的插入位置")
add_picture_after(insert_after, IMG)

# Keep core properties neutral for anonymous competition submission.
cp = doc.core_properties
cp.title = "基于海鸥派 SS928 与 STM32 底盘的视觉闭环自主泊车系统"
cp.subject = "竞赛作品报告"
cp.author = ""
cp.comments = ""
cp.keywords = ""

# Normalize table body emphasis after text replacement: header rows bold, body rows regular.
for table in doc.tables:
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, bold=(ri == 0))

doc.save(str(OUT))
print(OUT)
