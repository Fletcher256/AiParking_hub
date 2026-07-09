from docx import Document
from docx.shared import Inches
from pathlib import Path
p=Path(r'D:\parking_board_agent\tmp_image_doc.docx')
d=Document(); d.add_paragraph('before'); d.add_picture(r'D:\parking_board_agent\docs\line_follow_flowchart.png', width=Inches(5.65)); d.add_paragraph('after'); d.save(p)
print(p)
