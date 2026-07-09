from pathlib import Path
from PIL import Image, ImageOps, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

SRC_CANDIDATES = [
    Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版_补实物照片_补泊车前后对比.docx"),
    Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版_补实物照片.docx"),
    Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版.docx"),
]
SRC = next((p for p in SRC_CANDIDATES if p.exists()), None)
if SRC is None:
    raise FileNotFoundError("未找到可用的竞赛报告 DOCX")
OUT = Path(r"C:/Users/Cheng/Desktop/嵌入式大赛作品报告_视觉闭环自主泊车系统_竞赛提交版_补实物照片_补泊车前后对比_补YOLO识别效果.docx")
YOLO_IMG = Path(r"C:/Users/Cheng/Documents/Tencent Files/1571186517/nt_qq/nt_data/Pic/2026-07/Ori/8c8713a854afce0267cb1900cb7f409c.png")
FIG = Path(r"D:/parking_board_agent/docs/yolo_detection_effect_figure.jpg")

if not YOLO_IMG.exists():
    raise FileNotFoundError(YOLO_IMG)


def get_font(size=34, bold=False):
    candidates = [
        r"C:/Windows/Fonts/msyhbd.ttc" if bold else r"C:/Windows/Fonts/msyh.ttc",
        r"C:/Windows/Fonts/simhei.ttf",
        r"C:/Windows/Fonts/simsun.ttc",
        r"C:/Windows/Fonts/arial.ttf",
    ]
    for c in candidates:
        try:
            return ImageFont.truetype(c, size)
        except Exception:
            pass
    return ImageFont.load_default()

FONT_TITLE = get_font(48, True)
FONT_SUB = get_font(29, False)
FONT_LABEL = get_font(30, True)
FONT_NOTE = get_font(25, False)


def fit_contain(img, size):
    img = ImageOps.exif_transpose(img).convert("RGB")
    w, h = img.size
    tw, th = size
    scale = min(tw / w, th / h)
    nw, nh = int(w * scale + 0.5), int(h * scale + 0.5)
    resized = img.resize((nw, nh), Image.Resampling.LANCZOS)
    bg = Image.new("RGB", size, (247, 249, 252))
    bg.paste(resized, ((tw - nw) // 2, (th - nh) // 2))
    return bg


def make_yolo_figure():
    FIG.parent.mkdir(parents=True, exist_ok=True)
    src = Image.open(YOLO_IMG)
    W, H = 2400, 1320
    margin = 70
    header_h = 125
    footer_h = 92
    canvas = Image.new("RGB", (W, H), (255, 255, 255))
    d = ImageDraw.Draw(canvas)
    d.text((margin, 28), "YOLO 车位识别效果图", fill=(22, 44, 72), font=FONT_TITLE)
    d.text((margin, 83), "黄色 polygon 为板端识别/叠加的车位区域，蓝色点用于提示目标区域中心参考", fill=(82, 96, 116), font=FONT_SUB)

    box_x, box_y = margin, header_h + 16
    box_w, box_h = W - 2 * margin, H - header_h - footer_h - 24
    d.rounded_rectangle((box_x, box_y, box_x + box_w, box_y + box_h), radius=22, fill=(247, 249, 252), outline=(239, 188, 34), width=5)
    im = fit_contain(src, (box_w - 34, box_h - 34))
    canvas.paste(im, (box_x + 17, box_y + 17))

    # Small evidence labels outside the image area, unobtrusive.
    y = H - footer_h + 10
    d.rounded_rectangle((margin, y, margin + 410, y + 44), radius=12, fill=(255, 248, 219), outline=(239, 188, 34), width=2)
    d.text((margin + 18, y + 6), "识别对象：停车位区域", fill=(97, 73, 11), font=FONT_LABEL)
    d.rounded_rectangle((margin + 445, y, margin + 930, y + 44), radius=12, fill=(229, 246, 255), outline=(47, 148, 196), width=2)
    d.text((margin + 463, y + 6), "输出形式：polygon / overlay", fill=(24, 94, 130), font=FONT_LABEL)
    d.text((margin, H - 35), "注：该图为 YOLO 识别叠加效果截图，用于支撑车位感知与位姿估计输入链路。", fill=(82, 96, 116), font=FONT_NOTE)
    canvas.save(FIG, quality=92, optimize=True)


def set_run_font(run, name="微软雅黑", size=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_text(paragraph, text, size=None, bold=None, italic=None):
    for r in paragraph.runs:
        r.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    set_run_font(run, size=size, bold=bold, italic=italic)


def move_after(new_para, cursor_para):
    cursor_para._p.addnext(new_para._p)
    return new_para


def add_yolo_block_after(doc, after_para):
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(6)
    intro.paragraph_format.space_after = Pt(4)
    intro.paragraph_format.keep_with_next = True
    r = intro.add_run("图证补充：YOLO 识别效果图展示了系统对黄色胶带车位区域的检测与 polygon 叠加结果，该 polygon 后续作为 homography 位姿估计与闭环决策的输入。")
    set_run_font(r, size=10.5)
    cursor = move_after(intro, after_para)

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(3)
    pic_p.paragraph_format.space_after = Pt(2)
    pic_p.paragraph_format.keep_with_next = True
    run = pic_p.add_run()
    run.add_picture(str(FIG), width=Inches(6.25))
    for drawing in run._element.xpath(".//w:drawing"):
        doc_pr = drawing.xpath(".//wp:docPr")
        if doc_pr:
            doc_pr[0].set("name", "YOLO 车位识别效果图")
            doc_pr[0].set("descr", "YOLO overlay 截图：黄色 polygon 标出停车位区域，蓝色点提示目标区域中心参考。")
    cursor = move_after(pic_p, cursor)

    cap = doc.add_paragraph()
    try:
        cap.style = doc.styles["Caption"]
    except Exception:
        cap.style = doc.styles["Normal"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(9)
    set_paragraph_text(cap, "图4 YOLO 车位识别效果：黄色 polygon 标出识别到的停车位区域，识别结果用于后续 homography 位姿估计和闭环控制。", size=9, italic=True)
    move_after(cap, cursor)


make_yolo_figure()

doc = Document(str(SRC))

# Update evidence cells/references.
for table in doc.tables:
    for row in table.rows:
        joined = " | ".join(cell.text for cell in row.cells)
        if "YOLO 车位识别" in joined and len(row.cells) >= 3:
            p = row.cells[2].paragraphs[0]
            text = p.text.strip()
            if "图4" not in text:
                text = text.rstrip("。") + "；识别效果见图4。"
            set_paragraph_text(p, text, size=9)
        if row.cells and "感知与位姿" in row.cells[0].text and len(row.cells) >= 3:
            p = row.cells[2].paragraphs[0]
            set_paragraph_text(p, "YOLO 识别效果图、识别截图/视频、位姿 JSONL。", size=9)

# Insert after existing Figure 3 caption if available, keeping figure numbering sequential.
insert_after = None
for p in doc.paragraphs:
    if "图3 泊车前后对比" in p.text:
        insert_after = p
        break
if insert_after is None:
    for p in doc.paragraphs:
        if "表9 验证结果汇总" in p.text:
            insert_after = p
            break
if insert_after is None:
    for p in doc.paragraphs:
        if p.text.strip() == "2.3.2 软件各模块介绍":
            insert_after = p
            break
if insert_after is None:
    raise RuntimeError("未找到 YOLO 图插入位置")

add_yolo_block_after(doc, insert_after)

cp = doc.core_properties
cp.title = "基于海鸥派 SS928 与 STM32 底盘的视觉闭环自主泊车系统"
cp.subject = "竞赛作品报告"
cp.author = ""
cp.comments = ""
cp.keywords = ""

doc.save(str(OUT))
print(f"SRC={SRC}")
print(f"OUT={OUT}")
print(f"FIG={FIG}")
